from pathlib import Path

from docx import Document


ROOT = Path(r"D:\xilin")


def replace_paragraph_containing(document, needle, replacement):
    for paragraph in document.paragraphs:
        if needle in paragraph.text:
            paragraph.text = paragraph.text.replace(needle, replacement)
            return True
    return False


def append_section(document, heading, paragraphs):
    if any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        return
    document.add_paragraph(heading, style="Heading 1")
    for text in paragraphs:
        document.add_paragraph(text, style="Normal")


def update_product_document():
    path = ROOT / "托管班系统_第二版产品档案与业务流程方案.docx"
    document = Document(path)
    replace_paragraph_containing(
        document,
        "状态：基于第一版的流程深化稿，供开发前确认",
        "状态：第二版实施补充稿，家长自助入班与教师邀请审核流程已进入代码验证",
    )
    replace_paragraph_containing(
        document,
        "家长使用微信登录并通过绑定码、二维码或手机号校验绑定孩子。",
        "家长使用微信登录；优先通过教师邀请链接提交孩子信息，审核通过后自动建档和绑定，无需家长知道学生编号。",
    )
    append_section(
        document,
        "16. 家长自助入班与教师邀请流程（本次落地）",
        [
            "16.1 设计结论：家长不负责维护机构的学校、年级、班级基础数据，也不要求家长知道学生编号。教师从自己负责的学校班级发起邀请，家长只填写孩子姓名、家长手机号、称呼、关系和补充说明。",
            "16.2 邀请方式：教师在接送工作台选择负责班级后分享小程序页面，链接携带学校班级上下文。家长从该链接进入时，学校和班级自动带入并隐藏，降低“一年级2班、1年级二班、1年纪二班”等自由填写差异。正式二维码可以在后续版本替换链接实现，但不改变业务流程。",
            "16.3 无邀请兜底：家长仍可以直接填写学校和年级班级文本。系统先按机构内学校名称匹配，再对年级和班级文本做标准化识别；无法唯一匹配时不擅自建档，保留申请并交由管理端辅助归类。",
            "16.4 审核状态：申请状态为待审核、待补充资料、已通过、未通过。教师只看到自己被授权班级的申请，管理员可以查看全部申请；教师不能通过修改 URL 中的申请 ID 审核其他班级。",
            "16.5 通过规则：教师确认后，系统在目标学校班级内按姓名匹配学生。没有同名档案时自动创建正式学生档案；只有一条匹配时复用原档案；出现多条同名时返回待选择提示，不自动绑定错误孩子。通过后自动建立家长绑定，并在家长端生成入班通过通知。",
            "16.6 补充资料：教师可以退回“待补充资料”，家长端会显示审核备注，并可带入原申请修改后重新提交；重新提交回到待审核队列，不产生第二条重复申请。已存在同一家长、同一孩子姓名、同一学校班级的有效申请或绑定时，系统拒绝重复提交；被拒绝的申请允许重新发起。",
            "16.7 家长端体验：家长端同时展示申请记录和已绑定孩子。申请通过前只显示审核状态和备注，不展示未授权的接送、作业和照片；通过后自动进入孩子动态、接送、作业、请假和通知闭环。",
            "16.8 验收标准：真实走通教师分享、家长提交、教师退回补充、家长重提、教师通过、自动建档、自动绑定和家长收到通知；同时验证重复申请、班级文本标准化、同名学生和教师跨班越权均有明确结果。",
        ],
    )
    document.save(path)


def update_technical_document():
    path = ROOT / "托管班系统_第二版技术方案.docx"
    document = Document(path)
    replace_paragraph_containing(
        document,
        "状态：基于当前原型和第一版技术方案的技术深化稿",
        "状态：第二版实施补充稿，家长自助入班与教师邀请审核链路已实现并完成自动化验证",
    )
    replace_paragraph_containing(
        document,
        "开发环境可以保留手机号和学生档案编号的演示绑定；正式使用建议采用一次性绑定码或二维码，必要时由教师或主要监护人确认。绑定、解绑、邀请和主要监护人变更都要记录主体、时间和来源。",
        "家长使用微信登录后优先通过教师邀请上下文提交入班申请，不需要学生编号；开发环境仍保留旧绑定接口作为兼容入口。申请、审核、自动建档、绑定和通知都记录主体、时间、状态与来源；教师只能操作自己负责的学校班级。",
    )
    append_section(
        document,
        "17. 家长自助入班技术落地（本次确认）",
        [
            "17.1 数据模型：新增 parent_child_applications 表，保存 organization_id、parent_account_id、student_name、家长联系方式、家长原始学校/班级文本、解析后的 school_id/school_class_id、student_id、status、review_note、reviewer 和时间字段。状态使用 pending、needs_info、approved、rejected。",
            "17.2 家长接口：POST /api/v1/parent/child-applications 创建申请，GET /api/v1/parent/child-applications 查询当前家长申请，PUT /api/v1/parent/child-applications/{id} 只允许申请所属家长修改 needs_info 申请并重新提交。所有家长接口从 bearer Token 的 parent principal 获取身份，不信任请求体中的家长 ID。",
            "17.3 教师接口：GET /api/v1/child-applications 返回管理员全部申请或教师被授权班级的申请；POST /api/v1/child-applications/{id}/review 执行通过、退回补充和拒绝。教师审核前后都由服务端依据 TeacherAssignmentStore.FindByPair 校验 active 授权，前端隐藏按钮不构成安全边界。",
            "17.4 邀请上下文：小程序教师端分享 /pages/parent/index?schoolClassId={id}。家长端读取 schoolClassId 后直接提交学校班级 ID；服务端再次确认班级存在且 active，不把链接参数当作授权凭证。后续接入随机失效二维码时，只替换邀请解析层，保留申请表和审核 API。",
            "17.5 标准化与兜底：服务端统一去除空格、兼容“年纪/年级”、中文数字和括号写法，至少覆盖“一年级2班”“1年级2班”“1年纪二班”“一（2）班”。唯一匹配才回填学校班级；无法匹配时保留原始输入，避免把家长的近似文字直接写成错误正式档案。",
            "17.6 审核事务边界：审核通过时先在目标班级内按标准化姓名匹配 active 学生；零匹配创建 Student，单匹配复用，多匹配要求审核人进一步指定 student_id。随后写入 parent_student_bindings，并创建 child_application_approved 通知；绑定冲突按幂等成功处理。生产 MySQL 路径还应将建档、绑定、申请审核放入同一事务，避免半成品。",
            "17.7 重复与越权：服务层拦截同一家长、同一姓名、同一学校班级的 pending/needs_info/approved 申请或绑定；rejected 允许重新申请。教师列表过滤未授权班级，审核接口再次校验，跨班请求返回 403；家长查看申请时只按 parent_account_id 过滤，其他家长访问返回 404。",
            "17.8 前端落地：家长页取消学生编号输入，增加邀请班级提示、申请状态、审核备注和“修改并重提”；教师首页和接送页增加家长申请入口与班级邀请分享。已完成 TypeScript、ESLint、Stylelint、微信构建和 Web 构建验证。",
            "17.9 自动化验证：覆盖家长提交与查询、教师退回和家长重提、审核自动建档绑定、通过通知、重复申请、班级文本标准化、教师列表范围和跨班审核拒绝。接口契约同步写入 api/openapi.yaml，数据库迁移为 20260901001100_create_parent_child_applications.sql。",
        ],
    )
    document.save(path)


if __name__ == "__main__":
    update_product_document()
    update_technical_document()
